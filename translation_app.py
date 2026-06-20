#!/usr/bin/env python3
import json
import os
import re
import requests
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText
from typing import List, Dict
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def get_config_file_path() -> str:
    """获取配置文件的路径"""
    home_dir = os.path.expanduser("~")
    config_dir = os.path.join(home_dir, ".bidirectional_translator")
    os.makedirs(config_dir, exist_ok=True)
    return os.path.join(config_dir, "config.json")

def load_config() -> Dict:
    """加载配置文件"""
    config_file = get_config_file_path()
    if os.path.exists(config_file):
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return {}

def save_config(config: Dict):
    """保存配置到文件"""
    config_file = get_config_file_path()
    try:
        with open(config_file, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def count_chinese_chars(text: str) -> int:
    """计算中文文本的字符数（不包括空格和标点）"""
    count = 0
    for char in text:
        if '\u4e00' <= char <= '\u9fff':
            count += 1
    return count

def detect_language(text: str) -> str:
    """检测文本是中文还是英文"""
    text = text.strip()
    if not text:
        return "unknown"
    
    chinese_chars = 0
    english_chars = 0
    
    for char in text:
        if '\u4e00' <= char <= '\u9fff':
            chinese_chars += 1
        elif char.isalpha():
            english_chars += 1
    
    if chinese_chars > english_chars:
        return "zh"
    elif english_chars > chinese_chars:
        return "en"
    else:
        english_patterns = re.findall(r'[a-zA-Z]{2,}', text)
        if len(english_patterns) > 0:
            return "en"
        return "zh"

def count_chars(text: str, is_zh: bool) -> int:
    """计算文本长度"""
    if is_zh:
        return count_chinese_chars(text)
    else:
        return len(text.split())

def is_short_text(text: str, is_zh: bool) -> bool:
    """判断是否是短文本"""
    if is_zh:
        return count_chinese_chars(text) <= 20
    else:
        return len(text.split()) <= 10

def split_into_paragraphs(text: str, target_words: int = 80) -> List[str]:
    """分段函数，支持中英文"""
    lines = text.strip().split('\n')
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if re.match(r'^\s*[-—]+\s*$', line):
            continue
        line = re.sub(r'[-—]+', ' ', line)
        line = re.sub(r'\s+', ' ', line).strip()
        if line:
            cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)
    lang = detect_language(text)
    is_chinese = (lang == 'zh')
    
    paragraphs = []
    lines = text.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        
        is_title = False
        if re.match(r'^(Chapter|Section|Part|CHAPTER|SECTION|PART)\s+\d+', line, re.IGNORECASE):
            is_title = True
        elif re.match(r'^\d+(\.\d+)*\s+[A-Z]', line):
            is_title = True
        elif re.match(r'^(Summary|Conclusion|Introduction|Abstract|INTRODUCTION|SUMMARY|CONCLUSION|ABSTRACT)', line, re.IGNORECASE):
            is_title = True
        elif re.match(r'^第[一二三四五六七八九十百千万零〇\d]+[章节篇]', line):
            is_title = True
        elif re.match(r'^(摘要|引言|结论|总结|前言|序言)', line):
            is_title = True
        elif is_short_text(line, is_chinese):
            is_title = True
        
        if is_title:
            paragraphs.append(line)
            i += 1
            continue
        
        if is_chinese:
            sentences = re.split(r'(?<=[。！？])\s*', line)
        else:
            sentences = re.split(r'(?<=[.!?])\s+', line)
        
        current_paragraph = []
        current_length = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            sentence_length = count_chars(sentence, is_chinese)
            
            if sentence_length > target_words * 1.5:
                if current_paragraph:
                    paragraphs.append(' '.join(current_paragraph))
                    current_paragraph = []
                    current_length = 0
                paragraphs.append(sentence)
                continue
            
            if current_length + sentence_length <= target_words * 1.2:
                current_paragraph.append(sentence)
                current_length += sentence_length
            else:
                if current_paragraph:
                    paragraphs.append(' '.join(current_paragraph))
                current_paragraph = [sentence]
                current_length = sentence_length
        
        if current_paragraph:
            paragraphs.append(' '.join(current_paragraph))
        
        i += 1
    
    return paragraphs

def set_mixed_font(paragraph, text):
    """设置混合字体"""
    parts = []
    current_part = []
    current_is_chinese = None
    
    for char in text:
        is_chinese = '\u4e00' <= char <= '\u9fff'
        
        if current_is_chinese is None:
            current_is_chinese = is_chinese
            current_part.append(char)
        elif current_is_chinese == is_chinese:
            current_part.append(char)
        else:
            parts.append((''.join(current_part), current_is_chinese))
            current_is_chinese = is_chinese
            current_part = [char]
    
    if current_part:
        parts.append((''.join(current_part), current_is_chinese))
    
    for part_text, is_chinese in parts:
        run = paragraph.add_run(part_text)
        if is_chinese:
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            has_english_or_number = bool(re.search(r'[a-zA-Z0-9]', part_text))
            if has_english_or_number:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                run.font.name = 'SimSun'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(11)


class DeepSeekTranslator:
    def __init__(self, api_key: str, base_url: str = "https://api.deepseek.com", model: str = "deepseek-v4-pro"):
        self.api_key = api_key
        self.base_url = base_url
        self.model = model
        self.headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        self.default_system_prompt_en_to_zh = "你是一个专业的翻译专家，请将以下英文文本翻译成中文，保持原文的风格和语义。"
        self.default_system_prompt_zh_to_en = "You are a professional translation expert. Please translate the following Chinese text into English, maintaining the style and meaning of the original."
    
    def translate(self, text: str, source_lang: str = "auto", system_prompt: str = None) -> str:
        if system_prompt is None:
            if source_lang == "auto":
                source_lang = detect_language(text)
            
            if source_lang == "zh":
                system_prompt = self.default_system_prompt_zh_to_en
            else:
                system_prompt = self.default_system_prompt_en_to_zh
        
        payload = {
            "model": self.model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text}
            ],
            "temperature": 0.3
        }
        
        try:
            response = requests.post(
                f"{self.base_url}/v1/chat/completions",
                headers=self.headers,
                json=payload,
                timeout=60
            )
            response.raise_for_status()
            result = response.json()
            return result['choices'][0]['message']['content'].strip()
        except Exception as e:
            return f"翻译失败: {str(e)}"


class TranslationApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Bidirectional Translator v1.0.0")
        self.root.geometry("900x820")
        
        self.config = load_config()
        self.api_key = self.config.get("api_key", "")
        self.parallel_text = []
        self.current_file = ""
        self.translation_mode = "Default"
        self.custom_prompt = ""
        self.source_lang = "en"
        self.target_lang = "zh"
        self.is_paused = False
        self.is_stopped = False
        self.current_paragraphs = []
        self.current_index = 0
        self.system_prompt = None
        self.selected_model = self.config.get("selected_model", "deepseek-v4-pro")
        self.original_text = ""
        
        self.create_widgets()
    
    def create_widgets(self):
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        api_frame = ttk.LabelFrame(main_frame, text="API Configuration")
        api_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(api_frame, text="DeepSeek API Key:").pack(side=tk.LEFT, padx=5)
        self.api_entry = ttk.Entry(api_frame, width=60, show="*")
        self.api_entry.pack(side=tk.LEFT, padx=5)
        if self.api_key:
            self.api_entry.insert(0, self.api_key)
        else:
            self.api_entry.insert(0, os.environ.get("DEEPSEEK_API_KEY", ""))
        
        self.api_key_display = tk.Label(api_frame, text="", fg="gray", font=("Arial", 9))
        self.api_key_display.pack(side=tk.LEFT, padx=5)
        self.api_entry.bind('<KeyRelease>', self.update_api_key_display)
        self.api_entry.bind('<ButtonRelease-1>', self.update_api_key_display)
        self.update_api_key_display()
        
        model_frame = ttk.LabelFrame(main_frame, text="Model Selection")
        model_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(model_frame, text="DeepSeek Model:").pack(side=tk.LEFT, padx=5)
        
        self.model_options = {
            "DeepSeek V4 Pro": "deepseek-v4-pro",
            "DeepSeek V4 Flash": "deepseek-v4-flash"
        }
        default_model_display = "DeepSeek V4 Pro"
        for display_name, model_id in self.model_options.items():
            if model_id == self.selected_model:
                default_model_display = display_name
                break
        self.model_var = tk.StringVar(value=default_model_display)
        self.model_combobox = ttk.Combobox(
            model_frame, 
            textvariable=self.model_var,
            values=list(self.model_options.keys()),
            state="readonly",
            width=25
        )
        self.model_combobox.pack(side=tk.LEFT, padx=5)
        self.model_combobox.bind("<<ComboboxSelected>>", self.on_model_changed)
        
        mode_frame = ttk.LabelFrame(main_frame, text="Translation Mode")
        mode_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(mode_frame, text="Mode:").pack(side=tk.LEFT, padx=5)
        
        self.translation_mode_var = tk.StringVar(value="Default")
        mode_options = ["Default", "Translate with Prompt"]
        self.mode_combobox = ttk.Combobox(
            mode_frame, 
            textvariable=self.translation_mode_var,
            values=mode_options,
            state="readonly",
            width=25
        )
        self.mode_combobox.pack(side=tk.LEFT, padx=5)
        self.mode_combobox.bind("<<ComboboxSelected>>", self.on_mode_changed)
        
        self.prompt_frame = ttk.LabelFrame(main_frame, text="Custom Translation Prompt")
        self.prompt_visible = False
        
        ttk.Label(self.prompt_frame, text="Translation Prompt:").pack(anchor=tk.W, padx=5, pady=5)
        self.prompt_text = ScrolledText(self.prompt_frame, height=5, wrap=tk.WORD)
        self.prompt_text.pack(fill=tk.X, padx=5, pady=5)
        self.prompt_text.insert(
            tk.END, 
            "You are a professional translation expert. Please translate the following text into the target language, maintaining the style and meaning of the original."
        )
        
        control_frame = ttk.LabelFrame(main_frame, text="Actions")
        control_frame.pack(fill=tk.X, pady=5)
        
        self.open_btn = ttk.Button(control_frame, text="Open File", command=self.open_file)
        self.open_btn.pack(side=tk.LEFT, padx=5)
        
        self.split_btn = ttk.Button(control_frame, text="Split Text", command=self.split_text)
        self.split_btn.pack(side=tk.LEFT, padx=5)
        
        self.translate_btn = ttk.Button(control_frame, text="Translate", command=self.start_translation)
        self.translate_btn.pack(side=tk.LEFT, padx=5)
        
        self.export_btn = ttk.Button(control_frame, text="Export Word", command=self.export_to_word)
        self.export_btn.pack(side=tk.LEFT, padx=5)
        
        ttk.Label(control_frame, text="Words per Para:").pack(side=tk.LEFT, padx=5)
        self.words_entry = ttk.Entry(control_frame, width=5)
        self.words_entry.pack(side=tk.LEFT)
        self.words_entry.insert(0, "80")
        
        self.pause_btn = ttk.Button(control_frame, text="Pause", command=self.pause_translation, state=tk.DISABLED)
        self.pause_btn.pack(side=tk.LEFT, padx=10)
        
        text_frame = ttk.Frame(main_frame)
        text_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.left_frame = ttk.LabelFrame(text_frame, text="Source Text")
        self.left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
        
        self.input_text = ScrolledText(self.left_frame, wrap=tk.WORD, width=45)
        self.input_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.input_text.bind('<KeyRelease>', self.on_text_change)
        self.input_text.bind('<<Modified>>', self.on_input_changed)
        
        self.right_frame = ttk.LabelFrame(text_frame, text="Target Text")
        self.right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5)
        
        self.output_text = ScrolledText(self.right_frame, wrap=tk.WORD, width=45)
        self.output_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self.status_bar = ttk.Label(main_frame, text="Ready", relief=tk.SUNKEN)
        self.status_bar.pack(fill=tk.X, pady=5)
    
    def update_api_key_display(self, event=None):
        api_key = self.api_entry.get().strip()
        if api_key and len(api_key) >= 6:
            self.api_key_display.config(text=f"Key suffix: ...{api_key[-6:]}")
        else:
            self.api_key_display.config(text="")
        
        self.api_key = api_key
        self.config["api_key"] = api_key
        save_config(self.config)
    
    def on_mode_changed(self, event=None):
        mode = self.translation_mode_var.get()
        if mode == "Translate with Prompt":
            if not self.prompt_visible:
                self.prompt_frame.pack(fill=tk.X, pady=5, after=self.mode_combobox.master)
                self.prompt_visible = True
        else:
            if self.prompt_visible:
                self.prompt_frame.pack_forget()
                self.prompt_visible = False
    
    def on_model_changed(self, event=None):
        model_name = self.model_var.get()
        self.selected_model = self.model_options.get(model_name, "deepseek-v4-pro")
        self.status_bar.config(text=f"Selected model: {model_name}")
        self.config["selected_model"] = self.selected_model
        save_config(self.config)
    
    def on_text_change(self, event=None):
        pass
    
    def on_input_changed(self, event=None):
        if self.input_text.edit_modified():
            if not self.original_text:
                current_text = self.input_text.get(1.0, tk.END).strip()
                if current_text:
                    self.original_text = current_text
            self.input_text.edit_modified(False)
    
    def detect_input_language(self):
        text = self.input_text.get(1.0, tk.END).strip()
        if not text:
            messagebox.showwarning("Warning", "Please enter or open text first")
            return
        
        detected_lang = detect_language(text)
        is_prompt_mode = self.translation_mode_var.get() == "Translate with Prompt"
        
        if detected_lang == "zh":
            self.source_lang = "zh"
            self.target_lang = "en"
            self.left_frame.config(text="Source Text (Chinese)")
            self.right_frame.config(text="Target Text (English)")
            if not is_prompt_mode:
                self.prompt_text.delete(1.0, tk.END)
                self.prompt_text.insert(tk.END, "You are a professional translation expert. Please translate the following Chinese text into English, maintaining the style and meaning of the original.")
            self.status_bar.config(text="Detected: Chinese → English")
        else:
            self.source_lang = "en"
            self.target_lang = "zh"
            self.left_frame.config(text="Source Text (English)")
            self.right_frame.config(text="Target Text (Chinese)")
            if not is_prompt_mode:
                self.prompt_text.delete(1.0, tk.END)
                self.prompt_text.insert(tk.END, "你是一个专业的翻译专家，请将以下英文文本翻译成中文，保持原文的风格和语义。")
            self.status_bar.config(text="Detected: English → Chinese")
    
    def open_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt")])
        if file_path:
            self.current_file = file_path
            with open(file_path, 'r', encoding='utf-8') as f:
                file_content = f.read()
                self.original_text = file_content
                self.input_text.delete(1.0, tk.END)
                self.input_text.insert(tk.END, file_content)
            self.detect_input_language()
    
    def split_text(self):
        text_to_split = self.original_text.strip() if self.original_text and self.original_text.strip() else ""
        if not text_to_split:
            text_to_split = self.input_text.get(1.0, tk.END).strip()
        
        if not text_to_split:
            messagebox.showwarning("Warning", "Please enter or open text first")
            return
        
        try:
            target_words = int(self.words_entry.get())
            if target_words < 10:
                target_words = 10
                messagebox.showinfo("Info", "Minimum words per paragraph is 10")
            elif target_words > 200:
                target_words = 200
                messagebox.showinfo("Info", "Maximum words per paragraph is 200")
        except ValueError:
            target_words = 80
            messagebox.showwarning("Warning", "Please enter a valid number")
        
        paragraphs = split_into_paragraphs(text_to_split, target_words)
        self.input_text.delete(1.0, tk.END)
        self.input_text.insert(tk.END, '\n\n'.join(paragraphs))
        self.status_bar.config(text=f"Split into {len(paragraphs)} paragraphs")
    
    def start_translation(self):
        text_to_translate = self.original_text.strip() if self.original_text and self.original_text.strip() else self.input_text.get(1.0, tk.END).strip()
        if not text_to_translate:
            messagebox.showwarning("Warning", "Please enter or open text first")
            return
        
        self.api_key = self.api_entry.get().strip()
        if not self.api_key:
            messagebox.showwarning("Warning", "Please enter API key")
            return
        
        try:
            target_words = int(self.words_entry.get())
            if target_words < 10:
                target_words = 10
            elif target_words > 200:
                target_words = 200
        except ValueError:
            target_words = 80
        
        paragraphs = split_into_paragraphs(text_to_translate, target_words)
        if not paragraphs:
            messagebox.showwarning("Warning", "No content to translate")
            return
        
        mode = self.translation_mode_var.get()
        self.system_prompt = None
        if mode == "Translate with Prompt":
            self.system_prompt = self.prompt_text.get(1.0, tk.END).strip()
            if not self.system_prompt:
                messagebox.showwarning("Warning", "Please enter custom translation prompt")
                return
        
        self.detect_input_language()
        
        self.current_paragraphs = paragraphs
        self.current_index = 0
        self.is_paused = False
        self.is_stopped = False
        
        self.parallel_text = []
        self.output_text.delete(1.0, tk.END)
        
        self.translate_btn.config(state=tk.DISABLED)
        self.pause_btn.config(state=tk.NORMAL, text="Pause")
        
        self.translation_thread = threading.Thread(target=self.translate_threaded)
        self.translation_thread.daemon = True
        self.translation_thread.start()
    
    def translate_threaded(self):
        while True:
            if self.is_stopped:
                self.root.after(0, lambda: self.status_bar.config(text="Translation stopped"))
                self.root.after(0, self.reset_buttons)
                return
            
            if self.is_paused:
                self.root.after(0, lambda: self.status_bar.config(text=f"Paused at paragraph {self.current_index}/{len(self.current_paragraphs)}"))
                import time
                time.sleep(0.1)
                continue
            
            if self.current_index >= len(self.current_paragraphs):
                self.root.after(0, lambda: self.status_bar.config(text=f"Translation completed! {len(self.parallel_text)} paragraphs"))
                self.root.after(0, self.reset_buttons)
                return
            
            para = self.current_paragraphs[self.current_index]
            self.root.after(0, lambda idx=self.current_index, total=len(self.current_paragraphs): 
                self.status_bar.config(text=f"Translating paragraph {idx + 1}/{total} with {self.model_var.get()}...")
            )
            
            translator = DeepSeekTranslator(self.api_key, model=self.selected_model)
            translation = translator.translate(para, self.source_lang, self.system_prompt)
            
            if translation.startswith("翻译失败") or translation.startswith("Translation failed"):
                self.root.after(0, lambda err=translation: self.status_bar.config(text=f"Error: {err}"))
                self.root.after(0, lambda err=translation: messagebox.showerror("Translation Error", err))
                self.root.after(0, self.reset_buttons)
                return
            
            if self.source_lang == "en":
                self.parallel_text.append({
                    "source_text": para,
                    "target_text": translation,
                    "source_lang": "en",
                    "target_lang": "zh"
                })
            else:
                self.parallel_text.append({
                    "source_text": para,
                    "target_text": translation,
                    "source_lang": "zh",
                    "target_lang": "en"
                })
            
            def update_ui(trans=translation):
                self.output_text.insert(tk.END, f"{trans}\n\n")
                self.output_text.see(tk.END)
            
            self.root.after(0, update_ui)
            self.current_index += 1
    
    def pause_translation(self):
        if self.is_paused:
            self.is_paused = False
            self.pause_btn.config(text="Pause")
            self.status_bar.config(text="Resuming...")
        else:
            self.is_paused = True
            self.pause_btn.config(text="Resume")
            self.status_bar.config(text=f"Paused at paragraph {self.current_index}/{len(self.current_paragraphs)}")
    
    def reset_buttons(self):
        self.translate_btn.config(state=tk.NORMAL)
        self.pause_btn.config(state=tk.DISABLED, text="Pause")
    
    def export_to_word(self):
        """导出翻译结果为Word文档"""
        try:
            if not self.parallel_text:
                messagebox.showwarning("Warning", "No translated content to export!")
                self.status_bar.config(text="Ready")
                return
            
            self.status_bar.config(text="Opening save dialog...")
            self.root.update()
            
            try:
                home_dir = os.path.expanduser("~")
                export_dir = os.path.join(home_dir, "Documents")
                if not os.path.exists(export_dir):
                    export_dir = os.path.join(home_dir, "Desktop")
                if not os.path.exists(export_dir):
                    export_dir = home_dir
            except:
                export_dir = None
            
            default_filename = "parallel_translation.docx"
            if self.current_file:
                default_filename = os.path.splitext(os.path.basename(self.current_file))[0] + "_parallel.docx"
            
            file_options = {
                "defaultextension": ".docx",
                "filetypes": [("Word Documents", "*.docx")],
                "initialfile": default_filename,
                "title": "Save Translation as Word Document"
            }
            if export_dir and os.path.exists(export_dir):
                file_options["initialdir"] = export_dir
            
            output_path = filedialog.asksaveasfilename(**file_options)
            
            if not output_path:
                self.status_bar.config(text="Export cancelled")
                return
            
            self.status_bar.config(text="Creating document...")
            self.root.update()
            
            doc = Document()
            
            if self.source_lang == "en":
                doc.add_heading('English-Chinese Parallel Translation', level=1)
            else:
                doc.add_heading('Chinese-English Parallel Translation', level=1)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            headers = table.rows[0].cells
            if self.source_lang == "en":
                headers[0].text = "English"
                headers[1].text = "Chinese"
            else:
                headers[0].text = "Chinese"
                headers[1].text = "English"
            
            total = len(self.parallel_text)
            for idx, item in enumerate(self.parallel_text):
                self.status_bar.config(text=f"Adding paragraph {idx + 1}/{total}...")
                self.root.update()
                
                row_cells = table.add_row().cells
                
                source_text = item['source_text'].strip()
                source_text = re.sub(r'[-—]+', ' ', source_text)
                source_text = re.sub(r'\n\s*\n', '\n', source_text)
                source_text = re.sub(r'\s+', ' ', source_text).strip()
                if source_text.startswith(('---', '-', '——', '—')):
                    source_text = re.sub(r'^[-—]+', '', source_text).strip()
                
                target_text = item['target_text'].strip()
                target_text = re.sub(r'[-—]+', ' ', target_text)
                target_text = re.sub(r'\n\s*\n', '\n', target_text)
                target_text = re.sub(r'\s+', ' ', target_text).strip()
                if target_text.startswith(('---', '-', '——', '—')):
                    target_text = re.sub(r'^[-—]+', '', target_text).strip()
                
                para_source = row_cells[0].paragraphs[0]
                para_source.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para_source.clear()
                if source_text:
                    set_mixed_font(para_source, source_text)
                
                para_target = row_cells[1].paragraphs[0]
                para_target.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para_target.clear()
                if target_text:
                    set_mixed_font(para_target, target_text)
            
            self.status_bar.config(text="Formatting table...")
            self.root.update()
            
            for row in table.rows:
                for cell in row.cells:
                    cell.width = Inches(3.5)
            
            self.status_bar.config(text="Saving document...")
            self.root.update()
            
            doc.save(output_path)
            
            self.status_bar.config(text="Export complete!")
            messagebox.showinfo("Success", f"File successfully exported to:\n{output_path}")
            
        except Exception as e:
            error_msg = f"Export failed: {str(e)}"
            self.status_bar.config(text=error_msg)
            print(f"Export Error: {e}")
            friendly_msg = str(e)
            if "Read-only file system" in friendly_msg:
                friendly_msg = "请选择您的文档或桌面文件夹作为保存位置"
            messagebox.showerror("Export Error", f"Failed to export file:\n{friendly_msg}")


if __name__ == "__main__":
    root = tk.Tk()
    app = TranslationApp(root)
    root.mainloop()
